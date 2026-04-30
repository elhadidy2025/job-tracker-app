import os
import re
import base64
import zipfile
from io import BytesIO
from datetime import date

os.environ["STREAMLIT_SERVER_FILE_WATCHER_TYPE"] = "none"
os.environ["TOKENIZERS_PARALLELISM"] = "false"

import pandas as pd
import plotly.express as px
import streamlit as st
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

try:
    import pdfplumber
except Exception:
    pdfplumber = None

try:
    import docx
except Exception:
    docx = None


# =========================================================
# CONFIG
# =========================================================
st.set_page_config(
    page_title="JobTracker Dashboard",
    page_icon="💼",
    layout="wide",
    initial_sidebar_state="expanded",
)

DATA_DIR = "data"
UPLOAD_DIR = os.path.join(DATA_DIR, "documents_uploads")

DATA_PATH = os.path.join(DATA_DIR, "applications.csv")
DOCUMENTS_PATH = os.path.join(DATA_DIR, "documents.csv")
NOTES_PATH = os.path.join(DATA_DIR, "notes.csv")

STATUS_OPTIONS = [
    "Applied",
    "Screening",
    "Phone Screen",
    "Interview",
    "Offer",
    "Rejected",
]

SOURCE_OPTIONS = [
    "LinkedIn",
    "Company Website",
    "Referral",
    "StepStone",
    "Indeed",
    "Other",
    "Email",
    "Glassdoor",
]

DOCUMENT_TYPES = [
    "CV",
    "Motivation Letter",
    "Cover Letter",
    "Certificate",
    "Transcript",
    "Portfolio",
    "Job Description",
    "Other",
]

APP_COLUMNS = [
    "Company",
    "Role",
    "Date Applied",
    "Status",
    "Source",
    "Location",
    "Next Step",
]

DOC_COLUMNS = [
    "Document Name",
    "Type",
    "Linked Application",
    "Company",
    "Role",
    "Date Added",
    "File Name",
    "File Path",
    "Notes",
]

NOTE_COLUMNS = [
    "Title",
    "Company",
    "Date Added",
    "Note",
]

STATUS_COLORS = {
    "Applied": "#3367D6",
    "Screening": "#5F98E5",
    "Phone Screen": "#3FA9F5",
    "Interview": "#8A5CF6",
    "Offer": "#31C48D",
    "Rejected": "#F1416C",
}

COMMON_SKILLS = [
    "python", "sql", "excel", "power bi", "tableau", "r", "pandas", "numpy",
    "scikit-learn", "machine learning", "deep learning", "statistics",
    "data analysis", "data visualization", "dashboard", "dashboards", "reporting",
    "business intelligence", "bi", "etl", "data warehouse", "spark",
    "cloud", "aws", "azure", "gcp", "docker", "git", "github",
    "nlp", "tensorflow", "pytorch", "streamlit", "plotly", "matplotlib",
    "database", "databases", "postgresql", "mysql", "snowflake", "api",
    "regression", "classification", "clustering", "time series",
    "data cleaning", "data modeling", "kpi", "kpis", "analytics",
    "forecasting", "optimization", "requirements", "documentation",
    "agile", "scrum", "english", "german",
    "deutsch", "englisch", "datenanalyse", "datenvisualisierung",
    "statistik", "praktikum", "werkstudent", "werkstudentin",
    "trainee", "vollzeit", "teilzeit", "remote", "hybrid",
    "berufserfahrung", "kommunikation", "präsentation",
]


# =========================================================
# STYLE
# =========================================================
st.markdown(
    """
<style>
.block-container {
    padding-top: 1rem;
    padding-left: 1.2rem;
    padding-right: 1.2rem;
    padding-bottom: 1rem;
    max-width: 100%;
}

[data-testid="stSidebar"] {
    background-color: #F7F8FA;
    border-right: 1px solid #E5E7EB;
}

h1 {
    color: #0F172A !important;
    font-size: 34px !important;
    font-weight: 850 !important;
    line-height: 1.2 !important;
}

h2, h3 {
    color: #0F172A !important;
    font-weight: 800 !important;
}

.stButton button {
    border-radius: 12px;
    font-weight: 700;
}

div[data-testid="stDataFrame"] {
    border-radius: 14px;
    overflow: hidden;
    border: 1px solid #E5E7EB;
}

.kpi-card {
    background: white;
    border: 1px solid #E5E7EB;
    border-radius: 18px;
    padding: 16px 16px 14px 16px;
    min-height: 125px;
    box-shadow: 0 2px 10px rgba(15, 23, 42, 0.04);
}

.kpi-top {
    display: flex;
    align-items: flex-start;
    gap: 12px;
}

.kpi-icon {
    width: 48px;
    height: 48px;
    border-radius: 16px;
    display: flex;
    align-items: center;
    justify-content: center;
    font-size: 24px;
    flex-shrink: 0;
}

.kpi-title {
    color: #475569;
    font-size: 14px;
    font-weight: 700;
    margin-bottom: 6px;
}

.kpi-value {
    color: #0F172A;
    font-size: 26px;
    font-weight: 850;
    line-height: 1;
    margin-bottom: 10px;
}

.kpi-sub {
    font-size: 13px;
    font-weight: 700;
}

.small-info-card {
    background: white;
    border: 1px solid #E5E7EB;
    border-radius: 16px;
    padding: 12px 14px;
    margin-bottom: 10px;
}

.small-info-title {
    font-size: 14px;
    font-weight: 800;
    color: #111827;
}

.small-info-sub {
    font-size: 13px;
    color: #64748B;
    margin-top: 4px;
}

.small-info-num {
    float: right;
    font-size: 22px;
    font-weight: 850;
    color: #8B5CF6;
}
</style>
""",
    unsafe_allow_html=True,
)


# =========================================================
# GENERAL HELPERS
# =========================================================
def stretch():
    return {"width": "stretch"}


def safe_filename(name):
    name = os.path.basename(str(name))
    name = re.sub(r"[^a-zA-Z0-9_. \-]", "_", name)
    return name.strip().replace(" ", "_")


def file_ext(path):
    return os.path.splitext(str(path).lower())[1]


def ensure_files():
    os.makedirs(DATA_DIR, exist_ok=True)
    os.makedirs(UPLOAD_DIR, exist_ok=True)

    if not os.path.exists(DATA_PATH):
        sample = pd.DataFrame(
            [
                ["Erste Group", "Machine Learning Intern", "2026-04-28", "Rejected", "LinkedIn", "Salzburg", "Archive application"],
                ["Infineon", "Product Data Analyst", "2026-04-25", "Screening", "Indeed", "Salzburg", "Send availability"],
                ["Magna", "Analytics Trainee", "2026-04-25", "Offer", "Company Website", "Graz", "Compare with other application"],
                ["Erste Group", "Business Intelligence Analyst", "2026-04-24", "Rejected", "Company Website", "Graz", "No action needed"],
                ["ams OSRAM", "Business Intelligence Analyst", "2026-04-20", "Interview", "LinkedIn", "Premstätten", "Upcoming interview"],
                ["AVL", "Data Analyst Working Student", "2026-04-18", "Phone Screen", "Referral", "Graz", "Prepare for recruiter call"],
            ],
            columns=APP_COLUMNS,
        )
        sample.to_csv(DATA_PATH, index=False)

    if not os.path.exists(DOCUMENTS_PATH):
        pd.DataFrame(columns=DOC_COLUMNS).to_csv(DOCUMENTS_PATH, index=False)

    if not os.path.exists(NOTES_PATH):
        pd.DataFrame(columns=NOTE_COLUMNS).to_csv(NOTES_PATH, index=False)


def normalize_status(value):
    if pd.isna(value) or str(value).strip() == "":
        return "Applied"

    raw = str(value).strip().lower()
    mapping = {
        "applied": "Applied",
        "sent": "Applied",
        "submitted": "Applied",
        "waiting": "Applied",
        "pending": "Applied",
        "screening": "Screening",
        "screen": "Screening",
        "phone screen": "Phone Screen",
        "phone": "Phone Screen",
        "call": "Phone Screen",
        "interview": "Interview",
        "interviewing": "Interview",
        "offer": "Offer",
        "accepted": "Offer",
        "rejected": "Rejected",
        "declined": "Rejected",
        "abgelehnt": "Rejected",
        "absage": "Rejected",
        "angebot": "Offer",
    }
    return mapping.get(raw, str(value).strip().title())


def normalize_source(value):
    if pd.isna(value) or str(value).strip() == "":
        return "Other"

    raw = str(value).strip().lower()
    mapping = {
        "linkedin": "LinkedIn",
        "company": "Company Website",
        "company website": "Company Website",
        "website": "Company Website",
        "career page": "Company Website",
        "referral": "Referral",
        "stepstone": "StepStone",
        "indeed": "Indeed",
        "glassdoor": "Glassdoor",
        "other": "Other",
        "email": "Email",
        "mail": "Email",
    }
    return mapping.get(raw, str(value).strip().title())


def dataframe_to_csv_bytes(dataframe):
    export_df = dataframe.copy()

    if "Date Applied" in export_df.columns:
        export_df["Date Applied"] = pd.to_datetime(
            export_df["Date Applied"],
            errors="coerce",
        ).dt.strftime("%Y-%m-%d")

    if "Date Added" in export_df.columns:
        export_df["Date Added"] = pd.to_datetime(
            export_df["Date Added"],
            errors="coerce",
        ).dt.strftime("%Y-%m-%d")

    return export_df.to_csv(index=False).encode("utf-8")


# =========================================================
# LOAD / SAVE
# =========================================================
def load_apps():
    ensure_files()
    df = pd.read_csv(DATA_PATH)

    for col in APP_COLUMNS:
        if col not in df.columns:
            df[col] = ""

    df = df[APP_COLUMNS].copy()
    df["Company"] = df["Company"].fillna("").astype(str).str.strip().replace("", "Unknown Company")
    df["Role"] = df["Role"].fillna("").astype(str).str.strip().replace("", "Unknown Role")
    df["Location"] = df["Location"].fillna("").astype(str).str.strip().replace("", "—")
    df["Next Step"] = df["Next Step"].fillna("").astype(str).str.strip().replace("", "Awaiting response")

    df["Status"] = df["Status"].apply(normalize_status)
    df["Source"] = df["Source"].apply(normalize_source)
    df["Date Applied"] = pd.to_datetime(df["Date Applied"], errors="coerce")
    df["Date Applied"] = df["Date Applied"].fillna(pd.Timestamp(date.today()))
    df["Month"] = df["Date Applied"].dt.strftime("%b %Y")

    return df


def save_apps(df):
    for col in APP_COLUMNS:
        if col not in df.columns:
            df[col] = ""

    df = df[APP_COLUMNS].copy()
    df["Date Applied"] = pd.to_datetime(df["Date Applied"], errors="coerce")
    df["Date Applied"] = df["Date Applied"].fillna(pd.Timestamp(date.today()))
    df["Date Applied"] = df["Date Applied"].dt.strftime("%Y-%m-%d")
    df.to_csv(DATA_PATH, index=False)


def load_docs():
    ensure_files()
    df = pd.read_csv(DOCUMENTS_PATH)

    for col in DOC_COLUMNS:
        if col not in df.columns:
            df[col] = ""

    df = df[DOC_COLUMNS].copy()
    df["Date Added"] = pd.to_datetime(df["Date Added"], errors="coerce")

    return df


def save_docs(df):
    for col in DOC_COLUMNS:
        if col not in df.columns:
            df[col] = ""

    df = df[DOC_COLUMNS].copy()
    df["Date Added"] = pd.to_datetime(df["Date Added"], errors="coerce").dt.strftime("%Y-%m-%d")
    df.to_csv(DOCUMENTS_PATH, index=False)


def load_notes():
    ensure_files()
    df = pd.read_csv(NOTES_PATH)

    for col in NOTE_COLUMNS:
        if col not in df.columns:
            df[col] = ""

    df = df[NOTE_COLUMNS].copy()
    df["Date Added"] = pd.to_datetime(df["Date Added"], errors="coerce")
    return df


def save_notes(df):
    for col in NOTE_COLUMNS:
        if col not in df.columns:
            df[col] = ""

    df = df[NOTE_COLUMNS].copy()
    df["Date Added"] = pd.to_datetime(df["Date Added"], errors="coerce").dt.strftime("%Y-%m-%d")
    df.to_csv(NOTES_PATH, index=False)


# =========================================================
# FILE HELPERS
# =========================================================
def save_uploaded_file(uploaded_file, doc_type, company, role):
    if uploaded_file is None:
        return "", ""

    original_name = safe_filename(uploaded_file.name)

    prefix = "_".join([
        safe_filename(doc_type or "Document"),
        safe_filename(company or "General"),
        safe_filename(role or "General"),
    ])

    final_name = f"{date.today().strftime('%Y%m%d')}_{prefix}_{original_name}"
    path = os.path.join(UPLOAD_DIR, final_name)

    with open(path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    return final_name, path


def read_pdf_path(path):
    if pdfplumber is None:
        return ""

    parts = []

    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                txt = page.extract_text()
                if txt:
                    parts.append(txt)
    except Exception:
        return ""

    return "\n".join(parts)


def read_docx_path(path):
    if docx is None:
        return ""

    try:
        document = docx.Document(path)
        return "\n".join([p.text for p in document.paragraphs])
    except Exception:
        return ""


def read_txt_path(path):
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception:
        return ""


def read_text_from_path(path):
    ext = file_ext(path)

    if ext == ".pdf":
        return read_pdf_path(path)

    if ext == ".docx":
        return read_docx_path(path)

    if ext == ".txt":
        return read_txt_path(path)

    return ""


def read_uploaded_cv(uploaded_file):
    if uploaded_file is None:
        return ""

    lower = uploaded_file.name.lower()

    try:
        if lower.endswith(".pdf") and pdfplumber is not None:
            parts = []
            with pdfplumber.open(uploaded_file) as pdf:
                for page in pdf.pages:
                    txt = page.extract_text()
                    if txt:
                        parts.append(txt)
            return "\n".join(parts)

        if lower.endswith(".docx") and docx is not None:
            document = docx.Document(uploaded_file)
            return "\n".join([p.text for p in document.paragraphs])

        if lower.endswith(".txt"):
            return uploaded_file.read().decode("utf-8", errors="ignore")

    except Exception:
        return ""

    return ""


def preview_document(path, name):
    if not path or not os.path.exists(path):
        st.info("No saved file available for this document.")
        return

    ext = file_ext(path)

    if ext == ".pdf":
        with open(path, "rb") as f:
            encoded = base64.b64encode(f.read()).decode("utf-8")

        st.markdown(
            f"""
            <iframe src="data:application/pdf;base64,{encoded}"
                    width="100%"
                    height="650"
                    type="application/pdf">
            </iframe>
            """,
            unsafe_allow_html=True,
        )

    elif ext in [".png", ".jpg", ".jpeg"]:
        st.image(path, caption=name, width="stretch")

    elif ext in [".txt", ".docx"]:
        text = read_text_from_path(path)

        if text.strip():
            st.text_area(
                "Document Preview",
                value=text[:8000],
                height=420,
                key=f"preview_{safe_filename(name)}",
            )
        else:
            st.warning("Could not extract text from this file.")

    else:
        st.info("Preview not supported for this file type. Please download it.")


def create_zip_for_docs(docs_df):
    buffer = BytesIO()
    count = 0

    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as zipf:
        for _, row in docs_df.iterrows():
            path = str(row.get("File Path", ""))
            fname = str(row.get("File Name", "")) or os.path.basename(path)

            if path and os.path.exists(path):
                linked = str(row.get("Linked Application", "")).strip()
                folder = "General"

                if linked and linked != "Not linked / General":
                    folder = safe_filename(linked.replace("|", "_"))[:100]

                zipf.write(path, arcname=f"{folder}/{fname}")
                count += 1

    buffer.seek(0)
    return buffer, count


# =========================================================
# LABEL HELPERS
# =========================================================
def app_label(index, row):
    date_text = row["Date Applied"].strftime("%Y-%m-%d") if pd.notna(row["Date Applied"]) else ""
    return f"{index} | {row['Company']} | {row['Role']} | {date_text} | {row['Status']}"


def app_options(apps_df):
    options = ["Not linked / General"]

    if apps_df.empty:
        return options

    for idx, row in apps_df.sort_values("Date Applied", ascending=False).iterrows():
        options.append(app_label(idx, row))

    return options


def parse_app_label(label):
    if not label or label == "Not linked / General":
        return "", "", ""

    parts = [p.strip() for p in str(label).split("|")]

    if len(parts) >= 4:
        return label, parts[1], parts[2]

    return label, "", ""


# =========================================================
# ML / NLP HELPERS
# =========================================================
def clean_text(text):
    text = str(text).lower()
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[^a-zA-ZäöüÄÖÜß0-9+#./\\- ]", " ", text)
    return text.strip()


def extract_skills(text):
    cleaned = clean_text(text)
    found = []

    for skill in COMMON_SKILLS:
        pattern = r"\b" + re.escape(skill.lower()) + r"\b"

        if re.search(pattern, cleaned):
            found.append(skill)

    return sorted(set(found))


def tfidf_similarity(text_a, text_b):
    if not str(text_a).strip() or not str(text_b).strip():
        return 0

    vectorizer = TfidfVectorizer(ngram_range=(1, 2))
    matrix = vectorizer.fit_transform([text_a, text_b])
    score = cosine_similarity(matrix[0:1], matrix[1:2])[0][0]

    return round(score * 100)


def keyword_score(cv_text, job_text):
    cv_skills = set(extract_skills(cv_text))
    job_skills = set(extract_skills(job_text))

    if not job_skills:
        return 0, [], []

    matched = sorted(cv_skills.intersection(job_skills))
    missing = sorted(job_skills.difference(cv_skills))
    score = round((len(matched) / len(job_skills)) * 100)

    return score, matched, missing


def section_score(cv_text):
    cleaned = clean_text(cv_text)

    sections = {
        "Education / Ausbildung": ["education", "university", "master", "bachelor", "degree", "ausbildung", "studium"],
        "Experience / Erfahrung": ["experience", "work experience", "internship", "employment", "erfahrung", "praktikum"],
        "Skills / Fähigkeiten": ["skills", "technical skills", "tools", "technologies", "fähigkeiten", "kenntnisse"],
        "Projects / Projekte": ["projects", "project", "projekte", "projekt"],
        "Contact / Kontakt": ["email", "phone", "linkedin", "github", "kontakt", "telefon"],
    }

    found = []

    for section, keywords in sections.items():
        if any(k in cleaned for k in keywords):
            found.append(section)

    missing = sorted(set(sections.keys()) - set(found))
    score = round((len(found) / len(sections)) * 100)

    return score, found, missing


def final_match_score(semantic, keyword, section):
    return round((0.55 * semantic) + (0.30 * keyword) + (0.15 * section))


def analyze_job_description(job_text):
    cleaned = clean_text(job_text)
    skills = extract_skills(job_text)

    seniority = "Not clear"

    if re.search(r"\binternship\b|\bintern\b|\bworking student\b|\btrainee\b|\bpraktikum\b|\bwerkstudent\b", cleaned):
        seniority = "Internship / Working Student / Trainee"

    elif re.search(r"\bjunior\b|\bentry level\b|\bgraduate\b|\beinsteiger\b", cleaned):
        seniority = "Junior / Entry-level"

    elif re.search(r"\bsenior\b|\blead\b|\b5\+ years\b|\b7\+ years\b", cleaned):
        seniority = "Senior"

    elif re.search(r"\b2\+ years\b|\b3\+ years\b|\bprofessional\b|\bberufserfahrung\b", cleaned):
        seniority = "Mid-level"

    red_flags = []

    if re.search(r"\bc1\b|\bfluent german\b|\bnative german\b|\bverhandlungssicheres deutsch\b", cleaned):
        red_flags.append("Strong German requirement")

    if re.search(r"\bfull-time\b|\bvollzeit\b", cleaned):
        red_flags.append("Full-time requirement")

    if re.search(r"\b5\+ years\b|\b7\+ years\b|\bsenior\b", cleaned):
        red_flags.append("High experience requirement")

    responsibilities = []
    verbs = [
        "analyze", "build", "develop", "design", "maintain", "report", "visualize",
        "communicate", "collaborate", "optimize", "automate",
        "analysieren", "entwickeln", "erstellen", "berichten", "visualisieren",
        "kommunizieren", "optimieren",
    ]

    for sentence in re.split(r"[.\n]", job_text):
        sent = sentence.strip()

        if len(sent) > 25 and any(v in clean_text(sent) for v in verbs):
            responsibilities.append(sent)

    summary = (
        "This role appears to focus on "
        + (", ".join(skills[:6]) if skills else "general job-related responsibilities")
        + "."
    )

    return {
        "summary": summary,
        "skills": skills,
        "seniority": seniority,
        "red_flags": red_flags,
        "responsibilities": responsibilities[:6],
    }


# =========================================================
# METRICS / RENDER
# =========================================================
def calculate_metrics(df):
    total = len(df)
    interviews = len(df[df["Status"] == "Interview"])
    offers = len(df[df["Status"] == "Offer"])
    rejected = len(df[df["Status"] == "Rejected"])
    active = len(df[df["Status"].isin(["Applied", "Screening", "Phone Screen", "Interview"])])
    responded = len(df[df["Status"].isin(["Screening", "Phone Screen", "Interview", "Offer", "Rejected"])])

    return {
        "total": total,
        "interviews": interviews,
        "offers": offers,
        "rejected": rejected,
        "active": active,
        "response_rate": round((responded / total) * 100) if total else 0,
        "offer_rate": round((offers / total) * 100) if total else 0,
    }


def render_kpi(col, title, value, subtitle, icon, bg_color, sub_color="#10B981"):
    with col:
        st.markdown(
            f"""
            <div class="kpi-card">
                <div class="kpi-top">
                    <div class="kpi-icon" style="background:{bg_color};">{icon}</div>
                    <div>
                        <div class="kpi-title">{title}</div>
                        <div class="kpi-value">{value}</div>
                    </div>
                </div>
                <div class="kpi-sub" style="color:{sub_color};">{subtitle}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )


def reset_dashboard_filters():
    st.session_state["dashboard_month"] = "All"
    st.session_state["dashboard_status"] = "All"
    st.session_state["dashboard_source"] = "All"
    st.session_state["dashboard_location"] = "All"


# =========================================================
# APP STATE
# =========================================================
ensure_files()

if "page" not in st.session_state:
    st.session_state.page = "Dashboard"

for key, default in [
    ("dashboard_month", "All"),
    ("dashboard_status", "All"),
    ("dashboard_source", "All"),
    ("dashboard_location", "All"),
    ("confirm_delete_application", None),
    ("confirm_delete_document", None),
    ("confirm_delete_note", None),
]:
    if key not in st.session_state:
        st.session_state[key] = default

apps = load_apps()
docs = load_docs()
notes = load_notes()


# =========================================================
# SIDEBAR
# =========================================================
with st.sidebar:
    st.markdown("## 💼 JobTracker")

    sidebar_pages = [
        ("Dashboard", "📊"),
        ("Applications", "📄"),
        ("Calendar", "📅"),
        ("Companies", "🏢"),
        ("Documents", "📁"),
        ("Analytics", "📈"),
        ("Insights", "✨"),
        ("Notes", "📝"),
    ]

    for page_name, icon in sidebar_pages:
        btn_type = "primary" if st.session_state.page == page_name else "secondary"

        if st.button(f"{icon} {page_name}", type=btn_type, key=f"nav_{page_name}", **stretch()):
            st.session_state.page = page_name
            st.rerun()

    st.divider()

    st.markdown("### Mohamed Elhadidy")
    st.caption("mohamed.elhadidy@student.tugraz.at")

    st.markdown("**📌 Goal**")
    st.caption("Data / BI / ML roles")

    sb1, sb2, sb3 = st.columns(3)
    sb1.metric("Apps", len(apps))
    sb2.metric("Active", len(apps[apps["Status"].isin(["Applied", "Screening", "Phone Screen", "Interview"])]))
    sb3.metric("Int.", len(apps[apps["Status"] == "Interview"]))

    st.divider()

    b1, b2 = st.columns(2)

    with b1:
        st.link_button("LinkedIn", "https://www.linkedin.com/in/elhadidy19/", **stretch())

    with b2:
        st.link_button("GitHub", "https://github.com/elhadidy2025", **stretch())

    st.caption("Demo app_Elhadidy · CSV upload supported")


# =========================================================
# REUSABLE FORMS
# =========================================================
def add_application_ui(expanded=False):
    with st.expander("➕ Add New Application", expanded=expanded):
        with st.form("add_application_form", clear_on_submit=True):
            c1, c2, c3 = st.columns(3)

            with c1:
                company = st.text_input("Company", key="add_company")
                status = st.selectbox("Status", STATUS_OPTIONS, key="add_status")

            with c2:
                role = st.text_input("Role", key="add_role")
                source = st.selectbox("Source", SOURCE_OPTIONS, key="add_source")

            with c3:
                applied_date = st.date_input("Date Applied", value=date.today(), key="add_date")
                location = st.text_input("Location", value="Graz", key="add_location")

            next_step = st.text_input("Next Step", value="Awaiting response", key="add_next_step")

            if st.form_submit_button("Add Application", **stretch()):
                if not company.strip() or not role.strip():
                    st.error("Please enter at least Company and Role.")
                else:
                    current = load_apps()
                    new_row = pd.DataFrame(
                        [[
                            company.strip(),
                            role.strip(),
                            applied_date.strftime("%Y-%m-%d"),
                            status,
                            source,
                            location.strip(),
                            next_step.strip(),
                        ]],
                        columns=APP_COLUMNS,
                    )
                    save_apps(pd.concat([current[APP_COLUMNS], new_row], ignore_index=True))
                    st.success("Application added successfully.")
                    st.rerun()


def upload_csv_ui(expanded=False):
    with st.expander("📤 Upload CSV and Visualize", expanded=expanded):
        uploaded_file = st.file_uploader("Upload applications CSV", type=["csv"], key="applications_csv_upload")

        if uploaded_file is None:
            st.info("CSV columns can be: Company, Role, Date Applied, Status, Source, Location, Next Step.")
            return

        try:
            uploaded = pd.read_csv(uploaded_file)
        except Exception as e:
            st.error(f"Could not read CSV: {e}")
            return

        st.dataframe(uploaded.head(10), hide_index=True, **stretch())

        imported = uploaded.copy()

        defaults = {
            "Company": "Unknown Company",
            "Role": "Unknown Role",
            "Date Applied": date.today().strftime("%Y-%m-%d"),
            "Status": "Applied",
            "Source": "Other",
            "Location": "—",
            "Next Step": "Awaiting response",
        }

        for col in APP_COLUMNS:
            if col not in imported.columns:
                imported[col] = defaults[col]

        imported = imported[APP_COLUMNS].copy()
        imported["Status"] = imported["Status"].apply(normalize_status)
        imported["Source"] = imported["Source"].apply(normalize_source)

        b1, b2 = st.columns(2)

        if b1.button("Replace Current Data", type="primary", key="replace_csv", **stretch()):
            save_apps(imported)
            st.success("Current data replaced.")
            st.rerun()

        if b2.button("Merge With Current Data", key="merge_csv", **stretch()):
            current = load_apps()
            save_apps(pd.concat([current[APP_COLUMNS], imported], ignore_index=True))
            st.success("Uploaded CSV merged.")
            st.rerun()


# =========================================================
# PAGES
# =========================================================
def dashboard_page():
    st.markdown("<h1>Job Application<br>Tracker Dashboard</h1>", unsafe_allow_html=True)

    months = ["All"] + sorted(apps["Month"].dropna().unique().tolist())
    statuses = ["All"] + sorted(apps["Status"].dropna().unique().tolist())
    sources = ["All"] + sorted(apps["Source"].dropna().unique().tolist())
    locations = ["All"] + sorted(apps["Location"].dropna().astype(str).unique().tolist())

    f1, f2, f3, f4, f5 = st.columns([1, 1, 1, 1, 0.55])

    with f1:
        st.selectbox(" ", months, key="dashboard_month", label_visibility="collapsed")

    with f2:
        st.selectbox(" ", statuses, key="dashboard_status", label_visibility="collapsed")

    with f3:
        st.selectbox(" ", sources, key="dashboard_source", label_visibility="collapsed")

    with f4:
        st.selectbox(" ", locations, key="dashboard_location", label_visibility="collapsed")

    with f5:
        if st.button("Reset", key="dashboard_reset", **stretch()):
            reset_dashboard_filters()
            st.rerun()

    add_application_ui(expanded=False)
    upload_csv_ui(expanded=False)

    data = apps.copy()

    if st.session_state["dashboard_month"] != "All":
        data = data[data["Month"] == st.session_state["dashboard_month"]]

    if st.session_state["dashboard_status"] != "All":
        data = data[data["Status"] == st.session_state["dashboard_status"]]

    if st.session_state["dashboard_source"] != "All":
        data = data[data["Source"] == st.session_state["dashboard_source"]]

    if st.session_state["dashboard_location"] != "All":
        data = data[data["Location"] == st.session_state["dashboard_location"]]

    m = calculate_metrics(data)

    c1, c2, c3, c4, c5 = st.columns(5)

    render_kpi(c1, "Total Applications", m["total"], "Live from your data", "📄", "#E8EFFC", "#12A150")
    render_kpi(c2, "Interviews", m["interviews"], "Interview stage", "👥", "#F0E7FF", "#12A150")
    render_kpi(c3, "Offers", m["offers"], f"Offer rate: {m['offer_rate']}%", "💼", "#E7F8EE", "#12A150")
    render_kpi(c4, "Rejections", m["rejected"], "Rejected applications", "❌", "#FDECEC", "#F1416C")
    render_kpi(c5, "Response Rate", f"{m['response_rate']}%", "Non-applied responses", "％", "#FFF5D7", "#12A150")

    st.markdown("<br>", unsafe_allow_html=True)

    g1, g2, g3 = st.columns([1.05, 1.15, 1.1])

    with g1:
        st.markdown("### Application Pipeline")

        status_counts = data["Status"].value_counts().reindex(STATUS_OPTIONS, fill_value=0).reset_index()
        status_counts.columns = ["Status", "Applications"]

        total = max(len(data), 1)
        status_counts["Label"] = (
            status_counts["Applications"].astype(str)
            + " ("
            + ((status_counts["Applications"] / total) * 100).round(0).astype(int).astype(str)
            + "%)"
        )

        fig = px.bar(
            status_counts,
            x="Applications",
            y="Status",
            orientation="h",
            color="Status",
            color_discrete_map=STATUS_COLORS,
            text="Label",
        )

        fig.update_layout(
            height=300,
            showlegend=False,
            margin=dict(l=10, r=10, t=10, b=10),
            xaxis_title="",
            yaxis_title="",
            plot_bgcolor="white",
            paper_bgcolor="white",
        )

        fig.update_traces(textposition="outside")
        st.plotly_chart(fig, config={"displayModeBar": False}, **stretch())
        st.caption(f"Conversion rate (Applied ➜ Offer): {m['offer_rate']}%")

    with g2:
        st.markdown("### Applications by Month")

        month_counts = (
            data.groupby(data["Date Applied"].dt.to_period("M"))
            .size()
            .reset_index(name="Applications")
        )

        if not month_counts.empty:
            month_counts["Month"] = month_counts["Date Applied"].dt.strftime("%b %Y")

            fig = px.bar(
                month_counts,
                x="Month",
                y="Applications",
                text="Applications",
            )

            fig.update_layout(
                height=300,
                margin=dict(l=10, r=10, t=10, b=10),
                plot_bgcolor="white",
                paper_bgcolor="white",
                xaxis_title="",
                yaxis_title="",
                showlegend=False,
            )

            st.plotly_chart(fig, config={"displayModeBar": False}, **stretch())
        else:
            st.info("No data available.")

    with g3:
        st.markdown("### Applications by Source")

        source_counts = data["Source"].value_counts().reset_index()
        source_counts.columns = ["Source", "Applications"]

        if not source_counts.empty:
            fig = px.pie(
                source_counts,
                names="Source",
                values="Applications",
                hole=0.62,
            )

            fig.update_layout(
                height=300,
                margin=dict(l=10, r=10, t=10, b=10),
                annotations=[
                    dict(
                        text=f"{len(data)}<br>Total",
                        x=0.5,
                        y=0.5,
                        font_size=16,
                        showarrow=False,
                    )
                ],
            )

            st.plotly_chart(fig, config={"displayModeBar": False}, **stretch())
        else:
            st.info("No data available.")

    st.markdown("<br>", unsafe_allow_html=True)

    lower_left, lower_right = st.columns([2.5, 1.1])

    with lower_left:
        st.markdown("### Recent Applications")

        recent = data.sort_values("Date Applied", ascending=False).copy()
        recent["Date Applied"] = recent["Date Applied"].dt.strftime("%b %d, %Y")

        st.dataframe(recent[APP_COLUMNS], hide_index=True, height=330, **stretch())

    with lower_right:
        st.markdown("### Insights & Next Steps ✨")

        interview_df = data[data["Status"].isin(["Interview", "Phone Screen"])].copy()
        pending_df = data[data["Status"].isin(["Applied", "Screening", "Phone Screen", "Interview"])].copy()

        upcoming_count = len(interview_df)
        pending_count = len(pending_df)

        top_interview = "No interview scheduled"

        if not interview_df.empty:
            first_row = interview_df.sort_values("Date Applied", ascending=False).iloc[0]
            top_interview = f"{first_row['Company']} – {first_row['Role']}"

        st.markdown(
            f"""
            <div class="small-info-card">
                <div class="small-info-title">📅 Upcoming Interview <span class="small-info-num">{upcoming_count}</span></div>
                <div class="small-info-sub">{top_interview}</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        st.markdown(
            f"""
            <div class="small-info-card">
                <div class="small-info-title">⏰ Pending Follow-ups <span class="small-info-num" style="color:#F59E0B;">{pending_count}</span></div>
                <div class="small-info-sub">Applications that may need follow-up</div>
            </div>
            """,
            unsafe_allow_html=True,
        )

        if not pending_df.empty:
            next_items = pending_df.sort_values("Date Applied", ascending=False).head(6)

            for _, row in next_items.iterrows():
                st.markdown(f"- **{row['Company']}** — {row['Role']}  \n  *{row['Next Step']}*")

    st.download_button(
        "⬇️ Download Current Dashboard Data as CSV",
        data=dataframe_to_csv_bytes(data[APP_COLUMNS]),
        file_name="dashboard_applications.csv",
        mime="text/csv",
        key="download_dashboard_csv",
        **stretch(),
    )


def applications_page():
    st.markdown("<h1>Applications</h1>", unsafe_allow_html=True)

    add_application_ui(expanded=True)
    upload_csv_ui(expanded=False)

    st.markdown("## Application Table")

    s1, s2, s3 = st.columns([1.3, 1, 1])

    search = s1.text_input("Search", key="applications_search")
    status_filter = s2.multiselect("Filter by Status", STATUS_OPTIONS, key="applications_status_filter")
    source_filter = s3.multiselect("Filter by Source", SOURCE_OPTIONS, key="applications_source_filter")

    data = apps.copy()

    if search.strip():
        token = search.lower()
        data = data[
            data[APP_COLUMNS]
            .astype(str)
            .apply(lambda row: row.str.lower().str.contains(token).any(), axis=1)
        ]

    if status_filter:
        data = data[data["Status"].isin(status_filter)]

    if source_filter:
        data = data[data["Source"].isin(source_filter)]

    display = data.sort_values("Date Applied", ascending=False).copy()
    display["Date Applied"] = display["Date Applied"].dt.strftime("%Y-%m-%d")

    st.dataframe(display[APP_COLUMNS], hide_index=True, height=350, **stretch())

    st.markdown("### Download Applications CSV")

    d1, d2 = st.columns(2)

    with d1:
        st.download_button(
            "⬇️ Download Filtered Applications",
            data=dataframe_to_csv_bytes(display[APP_COLUMNS]),
            file_name="filtered_applications.csv",
            mime="text/csv",
            key="download_filtered_applications",
            **stretch(),
        )

    with d2:
        st.download_button(
            "⬇️ Download All Applications",
            data=dataframe_to_csv_bytes(apps[APP_COLUMNS]),
            file_name="all_applications.csv",
            mime="text/csv",
            key="download_all_applications",
            **stretch(),
        )

    st.markdown("## Edit / Delete Application")

    if apps.empty:
        st.info("No applications available.")
        return

    options = [
        app_label(idx, row)
        for idx, row in apps.sort_values("Date Applied", ascending=False).iterrows()
    ]

    selected = st.selectbox("Select application", options, key="application_select_edit")
    selected_idx = int(selected.split("|")[0].strip())
    row = apps.loc[selected_idx]

    with st.form("edit_application_form"):
        c1, c2, c3 = st.columns(3)

        with c1:
            company = st.text_input("Company", value=row["Company"], key="edit_company")
            status = st.selectbox(
                "Status",
                STATUS_OPTIONS,
                index=STATUS_OPTIONS.index(row["Status"]) if row["Status"] in STATUS_OPTIONS else 0,
                key="edit_status",
            )

        with c2:
            role = st.text_input("Role", value=row["Role"], key="edit_role")
            source = st.selectbox(
                "Source",
                SOURCE_OPTIONS,
                index=SOURCE_OPTIONS.index(row["Source"]) if row["Source"] in SOURCE_OPTIONS else 0,
                key="edit_source",
            )

        with c3:
            applied_date = st.date_input("Date Applied", value=row["Date Applied"].date(), key="edit_date")
            location = st.text_input("Location", value=row["Location"], key="edit_location")

        next_step = st.text_input("Next Step", value=row["Next Step"], key="edit_next_step")

        b1, b2 = st.columns(2)

        save_clicked = b1.form_submit_button("Save Changes", **stretch())
        delete_clicked = b2.form_submit_button("Delete Record", **stretch())

        if save_clicked:
            current = load_apps()
            current.loc[selected_idx, APP_COLUMNS] = [
                company,
                role,
                applied_date.strftime("%Y-%m-%d"),
                status,
                source,
                location,
                next_step,
            ]
            save_apps(current)
            st.success("Application updated.")
            st.rerun()

        if delete_clicked:
            st.session_state.confirm_delete_application = selected_idx
            st.rerun()

    if st.session_state.confirm_delete_application == selected_idx:
        st.warning(f"Are you sure you want to delete this application: {row['Company']} — {row['Role']}?")

        c1, c2 = st.columns(2)

        with c1:
            if st.button("Yes, delete application", key="confirm_delete_application_yes", **stretch()):
                current = load_apps().drop(index=selected_idx).reset_index(drop=True)
                save_apps(current)
                st.session_state.confirm_delete_application = None
                st.success("Application deleted.")
                st.rerun()

        with c2:
            if st.button("Cancel", key="confirm_delete_application_cancel", **stretch()):
                st.session_state.confirm_delete_application = None
                st.rerun()

    st.markdown("## Linked Documents")

    linked = docs[docs["Linked Application"].astype(str) == selected].copy()

    if linked.empty:
        st.info("No documents linked to this application.")
    else:
        linked["Date Added"] = linked["Date Added"].dt.strftime("%Y-%m-%d")

        st.dataframe(
            linked[["Document Name", "Type", "Date Added", "File Name", "Notes"]],
            hide_index=True,
            height=230,
            **stretch(),
        )


def calendar_page():
    st.markdown("<h1>Calendar</h1>", unsafe_allow_html=True)
    st.caption("Simple timeline view based on Date Applied and current application stage.")

    if apps.empty:
        st.info("No applications available.")
        return

    c1, c2 = st.columns([1, 1])

    year_options = ["All"] + sorted(apps["Date Applied"].dt.year.dropna().astype(int).unique().tolist())
    month_options = ["All"] + list(range(1, 13))

    selected_year = c1.selectbox("Year", year_options, key="calendar_year")
    selected_month = c2.selectbox("Month", month_options, key="calendar_month")

    cal_df = apps.copy()

    if selected_year != "All":
        cal_df = cal_df[cal_df["Date Applied"].dt.year == selected_year]

    if selected_month != "All":
        cal_df = cal_df[cal_df["Date Applied"].dt.month == selected_month]

    cal_df = cal_df.sort_values("Date Applied", ascending=True)

    m1, m2, m3 = st.columns(3)
    m1.metric("Events", len(cal_df))
    m2.metric("Upcoming Interviews", len(cal_df[cal_df["Status"].isin(["Interview", "Phone Screen"])]))
    m3.metric("Active Applications", len(cal_df[cal_df["Status"].isin(["Applied", "Screening", "Phone Screen", "Interview"])]))

    if not cal_df.empty:
        timeline_df = cal_df.copy()
        timeline_df["Label"] = timeline_df["Company"] + " — " + timeline_df["Role"]

        fig = px.scatter(
            timeline_df,
            x="Date Applied",
            y="Company",
            color="Status",
            color_discrete_map=STATUS_COLORS,
            hover_name="Label",
            hover_data=["Role", "Location", "Source", "Next Step"],
            size_max=14,
        )

        fig.update_layout(
            height=420,
            margin=dict(l=10, r=10, t=20, b=10),
            plot_bgcolor="white",
            paper_bgcolor="white",
            xaxis_title="Date",
            yaxis_title="Company",
        )

        st.plotly_chart(fig, config={"displayModeBar": False}, **stretch())

        table_df = cal_df.copy()
        table_df["Date Applied"] = table_df["Date Applied"].dt.strftime("%Y-%m-%d")

        st.dataframe(
            table_df[["Date Applied", "Company", "Role", "Status", "Location", "Next Step"]],
            hide_index=True,
            height=330,
            **stretch(),
        )
    else:
        st.info("No data for the selected period.")


def companies_page():
    st.markdown("<h1>Companies</h1>", unsafe_allow_html=True)

    if apps.empty:
        st.info("No applications available.")
        return

    grouped = apps.groupby("Company").agg(
        Applications=("Company", "count"),
        Roles=("Role", pd.Series.nunique),
        Interviews=("Status", lambda s: int((s == "Interview").sum())),
        Offers=("Status", lambda s: int((s == "Offer").sum())),
        Rejections=("Status", lambda s: int((s == "Rejected").sum())),
        Latest_Application=("Date Applied", "max"),
        Location=("Location", lambda x: ", ".join(sorted(set(x.astype(str))))),
    ).reset_index()

    grouped["Active"] = grouped["Applications"] - grouped["Offers"] - grouped["Rejections"]
    grouped["Latest_Application"] = pd.to_datetime(grouped["Latest_Application"]).dt.strftime("%Y-%m-%d")
    grouped = grouped.sort_values(["Applications", "Interviews"], ascending=[False, False])

    c1, c2, c3 = st.columns(3)
    c1.metric("Unique Companies", apps["Company"].nunique())
    c2.metric("Total Roles Applied", apps["Role"].nunique())
    c3.metric("Companies With Interviews", grouped[grouped["Interviews"] > 0]["Company"].nunique())

    st.markdown("## Company Overview")
    st.dataframe(grouped, hide_index=True, height=360, **stretch())

    company_list = sorted(apps["Company"].unique().tolist())
    selected_company = st.selectbox("Select company", company_list, key="selected_company_page")

    company_apps = apps[apps["Company"] == selected_company].copy().sort_values("Date Applied", ascending=False)
    company_apps_display = company_apps.copy()
    company_apps_display["Date Applied"] = company_apps_display["Date Applied"].dt.strftime("%Y-%m-%d")

    st.markdown(f"## Details — {selected_company}")
    st.dataframe(company_apps_display[APP_COLUMNS], hide_index=True, height=260, **stretch())

    linked_app_labels = [
        app_label(idx, row)
        for idx, row in company_apps.iterrows()
    ]

    company_docs = docs[docs["Linked Application"].isin(linked_app_labels)].copy()

    if company_docs.empty:
        st.info("No linked documents for this company.")
    else:
        company_docs["Date Added"] = company_docs["Date Added"].dt.strftime("%Y-%m-%d")

        st.dataframe(
            company_docs[["Document Name", "Type", "Date Added", "Role", "File Name", "Notes"]],
            hide_index=True,
            height=230,
            **stretch(),
        )


def documents_page():
    st.markdown("<h1>Documents</h1>", unsafe_allow_html=True)
    st.caption("Upload documents, link them to applications, preview them, and download them.")

    tab1, tab2, tab3 = st.tabs(["📁 Document Library", "🧾 Job Description Analyzer", "🎯 CV Match Score"])

    with tab1:
        with st.expander("➕ Upload / Add Document", expanded=True):
            with st.form("document_upload_form", clear_on_submit=True):
                selected_app = st.selectbox(
                    "Link to Application",
                    app_options(apps),
                    key="doc_link_application",
                )

                linked_app, linked_company, linked_role = parse_app_label(selected_app)

                c1, c2, c3 = st.columns(3)

                with c1:
                    uploaded = st.file_uploader(
                        "Upload Document",
                        type=["pdf", "docx", "txt", "png", "jpg", "jpeg"],
                        key="document_file_upload",
                    )

                    doc_name = st.text_input(
                        "Document Name",
                        placeholder="CV_Data_Science_V1.pdf",
                        key="document_name_input",
                    )

                    doc_type = st.selectbox(
                        "Type",
                        DOCUMENT_TYPES,
                        key="document_type_select",
                    )

                with c2:
                    company = st.text_input("Company", value=linked_company, key="document_company_input")
                    role = st.text_input("Role", value=linked_role, key="document_role_input")

                with c3:
                    added_date = st.date_input("Date Added", value=date.today(), key="document_date_input")

                notes_text = st.text_area("Notes", key="document_notes_text")

                if st.form_submit_button("Save Document", **stretch()):
                    if uploaded is None and not doc_name.strip():
                        st.error("Please upload a document or enter a document name.")
                    else:
                        if uploaded is not None and not doc_name.strip():
                            doc_name = uploaded.name

                        saved_name, saved_path = save_uploaded_file(
                            uploaded,
                            doc_type,
                            company.strip(),
                            role.strip(),
                        )

                        new_doc = pd.DataFrame(
                            [[
                                doc_name.strip(),
                                doc_type,
                                linked_app,
                                company.strip(),
                                role.strip(),
                                added_date.strftime("%Y-%m-%d"),
                                saved_name,
                                saved_path,
                                notes_text.strip(),
                            ]],
                            columns=DOC_COLUMNS,
                        )

                        current = load_docs()
                        save_docs(pd.concat([current[DOC_COLUMNS], new_doc], ignore_index=True))
                        st.success("Document saved.")
                        st.rerun()

        st.markdown("## Document Library")

        uploaded_docs = docs[
            docs["File Path"]
            .astype(str)
            .apply(lambda p: bool(str(p).strip()) and os.path.exists(str(p)))
        ].copy()

        if not uploaded_docs.empty:
            zip_buffer, file_count = create_zip_for_docs(uploaded_docs)

            st.download_button(
                f"⬇️ Download All Documents as ZIP ({file_count} files)",
                data=zip_buffer,
                file_name="job_tracker_documents.zip",
                mime="application/zip",
                key="download_documents_zip",
                **stretch(),
            )
        else:
            st.info("No uploaded files available for bulk download.")

        if docs.empty:
            st.info("No documents added yet.")
        else:
            display = docs.copy()
            display["Date Added"] = display["Date Added"].dt.strftime("%Y-%m-%d")

            st.dataframe(
                display[[
                    "Document Name",
                    "Type",
                    "Linked Application",
                    "Company",
                    "Role",
                    "Date Added",
                    "File Name",
                    "Notes",
                ]],
                hide_index=True,
                height=320,
                **stretch(),
            )

            st.markdown("## View / Download / Delete")

            options = []

            for idx, row in docs.iterrows():
                date_text = row["Date Added"].strftime("%Y-%m-%d") if pd.notna(row["Date Added"]) else "No date"
                link_text = row["Linked Application"] if str(row["Linked Application"]).strip() else "General"
                options.append(f"{idx} | {row['Document Name']} | {row['Type']} | {date_text} | {link_text}")

            selected_doc = st.selectbox("Select document", options, key="document_actions_select")
            selected_idx = int(selected_doc.split("|")[0].strip())
            selected_row = docs.loc[selected_idx]

            file_path = str(selected_row["File Path"])
            file_name = str(selected_row["File Name"] or selected_row["Document Name"])

            b1, b2, b3 = st.columns(3)

            view_clicked = b1.button("View Document", key="view_document_button", **stretch())

            with b2:
                if file_path and os.path.exists(file_path):
                    with open(file_path, "rb") as f:
                        st.download_button(
                            "Download Document",
                            data=f.read(),
                            file_name=file_name,
                            key="download_document_button",
                            **stretch(),
                        )
                else:
                    st.info("No saved file.")

            delete_clicked = b3.button("Delete Document", key="delete_document_button", **stretch())

            if delete_clicked:
                st.session_state.confirm_delete_document = selected_idx
                st.rerun()

            if st.session_state.confirm_delete_document == selected_idx:
                st.warning(f"Are you sure you want to delete this document: {selected_row['Document Name']}?")

                c1, c2 = st.columns(2)

                with c1:
                    if st.button("Yes, delete document", key="confirm_delete_document_yes", **stretch()):
                        if file_path and os.path.exists(file_path):
                            os.remove(file_path)

                        current = load_docs().drop(index=selected_idx).reset_index(drop=True)
                        save_docs(current)

                        st.session_state.confirm_delete_document = None
                        st.success("Document deleted.")
                        st.rerun()

                with c2:
                    if st.button("Cancel", key="confirm_delete_document_cancel", **stretch()):
                        st.session_state.confirm_delete_document = None
                        st.rerun()

            if view_clicked:
                st.markdown("### Preview")
                preview_document(file_path, file_name)

        st.markdown("## Documents by Application")

        linked_docs = docs[docs["Linked Application"].astype(str).str.strip() != ""].copy()

        if linked_docs.empty:
            st.info("No linked documents yet.")
        else:
            summary = linked_docs.groupby("Linked Application").agg(
                Documents=("Document Name", "count"),
                Types=("Type", lambda x: ", ".join(sorted(set(x.astype(str))))),
            ).reset_index().sort_values("Documents", ascending=False)

            st.dataframe(summary, hide_index=True, height=260, **stretch())

    with tab2:
        st.markdown("## Job Description Analyzer")

        jd_text = st.text_area("Paste Job Description", height=260, key="jd_analyzer_text")

        if st.button("Analyze Job Description", key="analyze_jd_button", **stretch()):
            if not jd_text.strip():
                st.error("Please paste a job description first.")
            else:
                analysis = analyze_job_description(jd_text)

                c1, c2, c3 = st.columns(3)
                c1.metric("Detected Skills", len(analysis["skills"]))
                c2.metric("Seniority", analysis["seniority"])
                c3.metric("Red Flags", len(analysis["red_flags"]))

                st.markdown("### Simple Summary")
                st.write(analysis["summary"])

                st.markdown("### Skills")
                st.write(", ".join(analysis["skills"]) if analysis["skills"] else "No skills detected.")

                st.markdown("### Responsibilities")

                if analysis["responsibilities"]:
                    for item in analysis["responsibilities"]:
                        st.markdown(f"- {item}")
                else:
                    st.info("No clear responsibility sentences detected.")

                st.markdown("### Red Flags")

                if analysis["red_flags"]:
                    for flag in analysis["red_flags"]:
                        st.markdown(f"- ⚠️ {flag}")
                else:
                    st.success("No obvious red flags detected.")

    with tab3:
        st.markdown("## CV ↔ Job Match Score")
        st.caption("Upload a new CV or use an existing uploaded CV from the Document Library.")

        source = st.radio(
            "Choose CV source",
            ["Upload new CV", "Use existing uploaded CV"],
            horizontal=True,
            key="cv_source_radio",
        )

        cv_text = ""

        left, right = st.columns([1, 1.4])

        with left:
            if source == "Upload new CV":
                cv_file = st.file_uploader("Upload CV", type=["pdf", "docx", "txt"], key="cv_new_upload")

                if cv_file is not None:
                    cv_text = read_uploaded_cv(cv_file)

            else:
                cv_docs = docs[
                    (docs["Type"].astype(str).str.lower() == "cv")
                    & (docs["File Path"].astype(str).str.strip() != "")
                ].copy()

                cv_docs = cv_docs[
                    cv_docs["File Path"].apply(lambda p: os.path.exists(str(p)))
                ]

                if cv_docs.empty:
                    st.warning("No uploaded CV files found in Document Library.")
                else:
                    cv_options = [
                        f"{idx} | {row['Document Name']} | {row['Company']} | {row['Role']}"
                        for idx, row in cv_docs.iterrows()
                    ]

                    selected_cv = st.selectbox("Select existing CV", cv_options, key="existing_cv_select")
                    cv_idx = int(selected_cv.split("|")[0].strip())
                    cv_row = cv_docs.loc[cv_idx]
                    cv_path = str(cv_row["File Path"])
                    cv_text = read_text_from_path(cv_path)

                    if st.button("Preview Selected CV", key="preview_selected_cv", **stretch()):
                        preview_document(cv_path, str(cv_row["File Name"]))

        with right:
            job_text = st.text_area("Paste Job Description", height=260, key="cv_match_job_text")

        if st.button("Calculate Match Score", key="calculate_match_score", **stretch()):
            if not cv_text.strip():
                st.error("Please upload a readable CV or choose an existing readable CV.")
                return

            if not job_text.strip():
                st.error("Please paste the job description first.")
                return

            semantic = tfidf_similarity(cv_text, job_text)
            keyword, matched, missing = keyword_score(cv_text, job_text)
            section, found_sections, missing_sections = section_score(cv_text)
            final = final_match_score(semantic, keyword, section)

            c1, c2, c3, c4 = st.columns(4)
            c1.metric("Final Match Score", f"{final}%")
            c2.metric("Similarity Score", f"{semantic}%")
            c3.metric("Keyword Score", f"{keyword}%")
            c4.metric("CV Section Score", f"{section}%")

            if final >= 80:
                st.success("Strong match. Your CV seems well aligned with this job.")
            elif final >= 60:
                st.info("Moderate match. Your CV is relevant, but could be tailored better.")
            else:
                st.warning("Weak to moderate match. Consider tailoring your CV more strongly.")

            r1, r2 = st.columns(2)

            with r1:
                st.markdown("### Matched Keywords")
                st.write(", ".join(matched) if matched else "No strong keyword overlap detected.")

                st.markdown("### Found CV Sections")
                st.write(", ".join(found_sections) if found_sections else "No common CV sections detected.")

            with r2:
                st.markdown("### Missing Keywords")
                st.write(", ".join(missing) if missing else "No major missing keywords detected.")

                st.markdown("### Missing CV Sections")
                st.write(", ".join(missing_sections) if missing_sections else "All basic CV sections were detected.")

            with st.expander("Extracted CV Text Preview"):
                st.text(cv_text[:4000])


def analytics_page():
    st.markdown("<h1>Analytics</h1>", unsafe_allow_html=True)

    if apps.empty:
        st.info("No data available.")
        return

    m = calculate_metrics(apps)

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Total Applications", m["total"])
    c2.metric("Unique Companies", apps["Company"].nunique())
    c3.metric("Sources Used", apps["Source"].nunique())
    c4.metric("Active Applications", m["active"])

    st.markdown("<br>", unsafe_allow_html=True)

    g1, g2 = st.columns(2)

    with g1:
        status_counts = apps["Status"].value_counts().reindex(STATUS_OPTIONS, fill_value=0).reset_index()
        status_counts.columns = ["Status", "Applications"]

        fig = px.bar(
            status_counts,
            x="Status",
            y="Applications",
            text="Applications",
            color="Status",
            color_discrete_map=STATUS_COLORS,
        )

        fig.update_layout(
            height=360,
            showlegend=False,
            margin=dict(l=10, r=10, t=20, b=10),
            plot_bgcolor="white",
            paper_bgcolor="white",
        )

        st.plotly_chart(fig, config={"displayModeBar": False}, **stretch())

    with g2:
        source_status = pd.crosstab(apps["Source"], apps["Status"]).reset_index()
        st.dataframe(source_status, hide_index=True, height=360, **stretch())

    monthly = apps.groupby(apps["Date Applied"].dt.to_period("M")).size().reset_index(name="Applications")

    if not monthly.empty:
        monthly["Month"] = monthly["Date Applied"].dt.strftime("%b %Y")

        fig = px.line(monthly, x="Month", y="Applications", markers=True)

        fig.update_layout(
            height=350,
            margin=dict(l=10, r=10, t=20, b=10),
            plot_bgcolor="white",
            paper_bgcolor="white",
        )

        st.plotly_chart(fig, config={"displayModeBar": False}, **stretch())


def insights_page():
    st.markdown("<h1>Insights</h1>", unsafe_allow_html=True)

    if apps.empty:
        st.info("No data available.")
        return

    m = calculate_metrics(apps)

    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Total", m["total"])
    c2.metric("Response Rate", f"{m['response_rate']}%")
    c3.metric("Offer Rate", f"{m['offer_rate']}%")
    c4.metric("Active", m["active"])
    c5.metric("Docs", len(docs))

    st.markdown("## Actionable Applications")

    actionable = apps[
        apps["Next Step"]
        .fillna("")
        .str.lower()
        .str.contains(
            "follow|await|response|recruiter|interview|assignment|call|onsite|availability|prepare|archive",
            regex=True,
        )
    ].copy()

    if actionable.empty:
        st.info("No actionable applications found.")
    else:
        actionable["Date Applied"] = actionable["Date Applied"].dt.strftime("%Y-%m-%d")
        st.dataframe(actionable[APP_COLUMNS], hide_index=True, height=320, **stretch())

    st.markdown("## Linked Documents Overview")

    linked = docs[docs["Linked Application"].astype(str).str.strip() != ""].copy()

    if linked.empty:
        st.info("No linked documents yet.")
    else:
        summary = linked["Linked Application"].value_counts().reset_index()
        summary.columns = ["Linked Application", "Documents"]
        st.dataframe(summary, hide_index=True, height=260, **stretch())


def notes_page():
    st.markdown("<h1>Notes</h1>", unsafe_allow_html=True)

    with st.expander("➕ Add Note", expanded=True):
        with st.form("note_form", clear_on_submit=True):
            c1, c2, c3 = st.columns(3)

            with c1:
                title = st.text_input("Title", key="note_title")

            with c2:
                company = st.text_input("Company", key="note_company")

            with c3:
                added_date = st.date_input("Date Added", value=date.today(), key="note_date")

            note_text = st.text_area("Note", key="note_text_area")

            if st.form_submit_button("Add Note", **stretch()):
                if not title.strip() or not note_text.strip():
                    st.error("Please enter at least title and note.")
                else:
                    current = load_notes()

                    new_note = pd.DataFrame(
                        [[
                            title.strip(),
                            company.strip(),
                            added_date.strftime("%Y-%m-%d"),
                            note_text.strip(),
                        ]],
                        columns=NOTE_COLUMNS,
                    )

                    save_notes(pd.concat([current[NOTE_COLUMNS], new_note], ignore_index=True))
                    st.success("Note added.")
                    st.rerun()

    st.markdown("## Notes List")

    if notes.empty:
        st.info("No notes added yet.")
    else:
        display = notes.copy()
        display["Date Added"] = display["Date Added"].dt.strftime("%Y-%m-%d")

        st.dataframe(display, hide_index=True, height=340, **stretch())

        options = []

        for idx, row in notes.iterrows():
            date_text = row["Date Added"].strftime("%Y-%m-%d") if pd.notna(row["Date Added"]) else "No date"
            options.append(f"{idx} | {row['Title']} | {row['Company']} | {date_text}")

        selected = st.selectbox("Select note to delete", options, key="selected_note_delete")
        selected_idx = int(selected.split("|")[0].strip())

        if st.button("Delete Note", key="delete_note_btn", **stretch()):
            st.session_state.confirm_delete_note = selected_idx
            st.rerun()

        if st.session_state.confirm_delete_note == selected_idx:
            st.warning(f"Are you sure you want to delete this note: {notes.loc[selected_idx, 'Title']}?")

            c1, c2 = st.columns(2)

            with c1:
                if st.button("Yes, delete note", key="confirm_delete_note_yes", **stretch()):
                    current = load_notes().drop(index=selected_idx).reset_index(drop=True)
                    save_notes(current)

                    st.session_state.confirm_delete_note = None
                    st.success("Note deleted.")
                    st.rerun()

            with c2:
                if st.button("Cancel", key="confirm_delete_note_cancel", **stretch()):
                    st.session_state.confirm_delete_note = None
                    st.rerun()


# =========================================================
# ROUTER
# =========================================================
if st.session_state.page == "Dashboard":
    dashboard_page()

elif st.session_state.page == "Applications":
    applications_page()

elif st.session_state.page == "Calendar":
    calendar_page()

elif st.session_state.page == "Companies":
    companies_page()

elif st.session_state.page == "Documents":
    documents_page()

elif st.session_state.page == "Analytics":
    analytics_page()

elif st.session_state.page == "Insights":
    insights_page()

elif st.session_state.page == "Notes":
    notes_page()